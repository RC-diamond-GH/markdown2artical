import argparse
import traceback
import re
import subprocess
import tempfile
import os
import io
import json
from bs4 import BeautifulSoup
import markdown as md_parser # Renamed to avoid conflict with os.markdown

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.font import Font

# --- Constants ---
# Font Names
FONT_SONGTI = "宋体"
FONT_HEITI = "黑体"
FONT_KAITI = "楷体"
FONT_TIMES_NEW_ROMAN = "Times New Roman"

# Font Sizes (in Pt)
SIZE_THREE = 16  # 三号
SIZE_SMALL_THREE = 15 # 小三号
SIZE_FOUR = 14   # 四号
SIZE_SMALL_FOUR = 12  # 小四号
SIZE_FIVE = 10.5 # 五号
SIZE_SMALL_FIVE = 9   # 小五号

# Line Spacing
LINE_SPACING_1_25 = 1.25
LINE_SPACING_1_5 = 1.5
LINE_SPACING_FIXED_20PT = Pt(20)

# Margins
MARGIN_CM = 2.5

with open('./config.json', 'r', encoding='utf-8') as f:
    config = json.load(f)

# --- Helper Functions ---

def set_run_font(run, east_asia_font=FONT_SONGTI, ascii_font=FONT_TIMES_NEW_ROMAN, size_pt=SIZE_SMALL_FOUR, bold=False, italic=False):
    """Sets font properties for a run, handling East Asian and ASCII characters."""
    font = run.font
    font.name = ascii_font
    font.element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
    font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)

def set_paragraph_formatting(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT, 
                             line_spacing_rule=WD_LINE_SPACING.ONE_POINT_FIVE, line_spacing_val=LINE_SPACING_1_25,
                             space_before_pt=0, space_after_pt=0, 
                             first_line_indent_cm=None, keep_with_next=False, page_break_before=False,
                             fixed_line_height_pt=None):
    """Applies common paragraph formatting."""
    fmt = paragraph.paragraph_format
    fmt.alignment = alignment
    if fixed_line_height_pt:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        fmt.line_spacing = fixed_line_height_pt
    else:
        fmt.line_spacing_rule = line_spacing_rule # Default is WD_LINE_SPACING.MULTIPLE for values like 1.25
        if line_spacing_rule == WD_LINE_SPACING.MULTIPLE or line_spacing_val > 3: # Heuristic for Pt vs multiplier
             fmt.line_spacing = line_spacing_val
        else: # For single, 1.5, double
            fmt.line_spacing = line_spacing_val


    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if first_line_indent_cm:
        fmt.first_line_indent = Cm(first_line_indent_cm)
    else:
        fmt.first_line_indent = None # Explicitly remove if not needed
        
    fmt.keep_with_next = keep_with_next
    fmt.page_break_before = page_break_before

def add_styled_paragraph(doc, text_content, default_east_asia_font, default_ascii_font, default_size_pt,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, 
                         line_spacing_rule=WD_LINE_SPACING.MULTIPLE, line_spacing_val=LINE_SPACING_1_25,
                         space_before_pt=0, space_after_pt=0,
                         first_line_indent_cm=None, bold_default=False,
                         fixed_line_height_pt=None, is_heading=False, heading_font_override=None,
                         is_reference=False):
    """Adds a paragraph with mixed Chinese/English font handling."""
    p = doc.add_paragraph()
    set_paragraph_formatting(p, alignment, line_spacing_rule, line_spacing_val,
                             space_before_pt, space_after_pt, first_line_indent_cm,
                             fixed_line_height_pt=fixed_line_height_pt)
    
    if is_reference: # directly add the reference text
        p.add_run(text_content)
        return p

    # Regex to split by English/numbers/symbols and Chinese characters
    # This regex attempts to capture sequences of non-Chinese and sequences of Chinese
    parts = re.split(r'([a-zA-Z0-9\s!"#$%&\'()*+,-./:;<=>?@\[\\\]^_`{|}~]+)', text_content)
    
    for part in parts:
        if not part:
            continue
        
        is_ascii_part = bool(re.match(r'^[a-zA-Z0-9\s!"#$%&\'()*+,-./:;<=>?@\[\\\]^_`{|}~]+$', part))
        
        # Handle in-text citations like [^N] or [N]
        citation_match = re.match(r'(\[\^?)(\d+)(\])', part) # Matches [^N] or [N]
        ref_match_in_text = re.match(r'(\[)(\d+(?:,\d+)*(?:-\d+)*)(\])', part) # Matches [1], [1,2], [1-3]

        if citation_match:
            # For superscript citations like [^1] in text becoming [1] superscript
            run = p.add_run(f"[{citation_match.group(2)}]")
            set_run_font(run, default_east_asia_font, FONT_TIMES_NEW_ROMAN, default_size_pt, bold=bold_default)
            run.font.superscript = True
        elif ref_match_in_text and not is_heading : # Avoid superscripting numbers in headings like "1.1 Title"
             # For inline citations like [1,2-5]
            run = p.add_run(f"[{ref_match_in_text.group(2)}]")
            set_run_font(run, default_east_asia_font, FONT_TIMES_NEW_ROMAN, default_size_pt, bold=bold_default)
            # Check if it should be superscript based on context (e.g. if it was [^...] originally)
            # This part is tricky without more context from markdown parser.
            # For now, assume if it's in the format [N] it's inline, not superscript unless it was [^N]
        
        elif is_ascii_part:
            run = p.add_run(part)
            current_font = heading_font_override if is_heading and heading_font_override else default_ascii_font
            set_run_font(run, default_east_asia_font, current_font, default_size_pt, bold=bold_default)
        else: # Chinese part or mixed part not caught by simple ASCII
            # Further split this part if it's mixed and not caught by the main regex
            sub_parts = re.split(r'([a-zA-Z0-9!"#$%&\'()*+,-./:;<=>?@\[\\\]^_`{|}~]+)', part)
            for sub_part in sub_parts:
                if not sub_part:
                    continue
                if bool(re.match(r'^[a-zA-Z0-9!"#$%&\'()*+,-./:;<=>?@\[\\\]^_`{|}~]+$', sub_part)):
                    run = p.add_run(sub_part)
                    current_font = heading_font_override if is_heading and heading_font_override else default_ascii_font
                    set_run_font(run, default_east_asia_font, current_font, default_size_pt, bold=bold_default)
                else:
                    run = p.add_run(sub_part)
                    current_font = heading_font_override if is_heading and heading_font_override else default_east_asia_font
                    set_run_font(run, current_font, default_ascii_font, default_size_pt, bold=bold_default)
    return p

def add_heading(doc, text, level, chapter_num_str=""):
    """Adds and styles a heading."""
    # Clean up heading text from markdown (e.g. "1.2.3 My Title" -> "My Title")
    clean_text = re.sub(r'^[\d\.]+\s*', '', text).strip()
    numbered_text = text # Keep original numbering for display

    if level == 1: # Chapter: "第一章 XXX"
        p = doc.add_paragraph(style='Heading 1')
        set_paragraph_formatting(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_5, page_break_before=True, space_after_pt=12) # Add some space after
        # "第X章" part
        run_chap_num = p.add_run(numbered_text.split(' ')[0] + (' ' if ' ' in numbered_text else '  ')) # Ensure space
        set_run_font(run_chap_num, FONT_HEITI, FONT_TIMES_NEW_ROMAN, SIZE_THREE, bold=True)
        # Title part
        run_title = p.add_run(' '.join(numbered_text.split(' ')[1:]))
        set_run_font(run_title, FONT_HEITI, FONT_HEITI, SIZE_THREE, bold=True) # Use HEITI for title too
        return p, numbered_text # Return original numbered text for TOC
        
    elif level == 2: # Section: "1.1 XXX"
        p = doc.add_paragraph(style='Heading 2')
        set_paragraph_formatting(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_5, space_before_pt=12, space_after_pt=6)
        # Number part (e.g., "1.1")
        num_match = re.match(r'([\d\.]+)\s*(.*)', text)
        if num_match:
            num_part = num_match.group(1)
            title_part = num_match.group(2)
            run_num = p.add_run(num_part + " ") # Space after number
            set_run_font(run_num, FONT_HEITI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_THREE, bold=True)
            run_title = p.add_run(title_part)
            set_run_font(run_title, FONT_HEITI, FONT_HEITI, SIZE_SMALL_THREE, bold=True)
        else: # Fallback
            run = p.add_run(text)
            set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_SMALL_THREE, bold=True)
        return p, numbered_text

    elif level == 3: # Sub-section: "1.1.1 XXX"
        p = doc.add_paragraph(style='Heading 3')
        set_paragraph_formatting(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing_val=LINE_SPACING_1_5, space_before_pt=10, space_after_pt=5)
        num_match = re.match(r'([\d\.]+)\s*(.*)', text)
        if num_match:
            num_part = num_match.group(1)
            title_part = num_match.group(2)
            run_num = p.add_run(num_part + " ") # Space after number
            set_run_font(run_num, FONT_HEITI, FONT_TIMES_NEW_ROMAN, SIZE_FOUR, bold=True)
            run_title = p.add_run(title_part)
            set_run_font(run_title, FONT_HEITI, FONT_HEITI, SIZE_FOUR, bold=True)
        else: # Fallback
            run = p.add_run(text)
            set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_FOUR, bold=True)
        return p, numbered_text

    elif level == 4: # "1. XXX"
        p = doc.add_paragraph()
        set_paragraph_formatting(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing_val=LINE_SPACING_1_25, space_before_pt=5, space_after_pt=2)
        run = p.add_run(text) # Whole text is bold
        set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_SMALL_FOUR, bold=True) # Small four, Heiti
        return p, None # Not in TOC

    elif level == 5: # "(1) XXX"
        # These are typically inline or start a paragraph without special heading formatting beyond default text.
        # The prompt implies "小标题内序号用⑴、⑵、⑶……，其余层次序号依次用A、B、C……，a、b、c……"
        # Markdown H5 `##### (1) Text` will be handled as a paragraph.
        # We'll rely on the add_styled_paragraph for this, assuming it's just a regular paragraph start.
        # If it needs specific styling (e.g. bold number), it should be handled during paragraph processing.
        # For now, H5 will just become a normal paragraph.
        return add_styled_paragraph(doc, text, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                    alignment=WD_ALIGN_PARAGRAPH.LEFT, line_spacing_val=LINE_SPACING_1_25), None

def convert_mermaid_to_image_mmdc(mermaid_code, temp_dir, output_format='png'):
    """Converts Mermaid code to an image file using mmdc."""
    # Create a unique filename for the mermaid source
    temp_mmd_file = tempfile.NamedTemporaryFile(mode='w+', delete=False, suffix='.mmd', dir=temp_dir, encoding='utf-8')
    #temp_mmd_file = open('./temp.mmd', 'w', encoding='utf-8')
    temp_mmd_file.write(mermaid_code)
    temp_mmd_file.close() # Close it so mmdc can read it
    
    image_file_path = temp_mmd_file.name.replace('.mmd', f'.{output_format}')

    try:
        # Ensure mmdc is in PATH or provide full path if necessary
        # Increased width for better resolution, adjust as needed
        # Added timeout for mmdc in case it hangs
        toRun = ['mmdc.cmd', '-i', temp_mmd_file.name, '-o', image_file_path, '-w', '1200', '-H', '800', '--scale', '1.5']
        #print(mermaid_code)
        process = subprocess.run(
            toRun, 
            check=True, capture_output=True, text=True, encoding='utf-8', timeout=30
        )
        if process.stderr:
            print(f"Mermaid CLI warning/info: {process.stderr}")
        return image_file_path
    except subprocess.CalledProcessError as e:
        print(f"Error converting Mermaid diagram using mmdc: {e.stderr}")
        return None
    except subprocess.TimeoutExpired:
        print("Error: mmdc command timed out.")
        return None
    except FileNotFoundError:
        print("Error: 'mmdc' (Mermaid CLI) not found. Please install it and ensure it's in your PATH.")
        #traceback.print_exc()
        return None
    finally:
        if os.path.exists(temp_mmd_file.name):
            os.remove(temp_mmd_file.name)

def add_image_with_caption(doc, image_path_or_stream, caption_text, fig_num_text, image_width_cm=15):
    """Adds an image and its caption to the document."""
    try:
        if isinstance(image_path_or_stream, str) and not os.path.exists(image_path_or_stream):
            print(f"Warning: Image file not found: {image_path_or_stream}. Skipping image.")
            # Add placeholder for caption
            p_caption = doc.add_paragraph()
            set_paragraph_formatting(p_caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_25)
            run_fig_num = p_caption.add_run(fig_num_text + " ")
            set_run_font(run_fig_num, FONT_KAITI, FONT_TIMES_NEW_ROMAN, SIZE_FIVE) # Fig num Times
            run_caption_text = p_caption.add_run(caption_text)
            set_run_font(run_caption_text, FONT_KAITI, FONT_KAITI, SIZE_FIVE) # Caption Kaiti
            add_styled_paragraph(doc, "", FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, space_after_pt=6) # Empty line after
            return

        # Add image, centered
        # To center image, add it to a paragraph that is centered.
        p_img = doc.add_paragraph()
        set_paragraph_formatting(p_img, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p_img.add_run().add_picture(image_path_or_stream, width=Cm(image_width_cm))
    except Exception as e:
        print(f"Error adding image '{image_path_or_stream}': {e}. Skipping image.")
        traceback.print_exc()
         # Add placeholder for caption even if image fails
        p_caption = doc.add_paragraph()
        set_paragraph_formatting(p_caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_25)
        run_fig_num = p_caption.add_run(fig_num_text + " ")
        set_run_font(run_fig_num, FONT_KAITI, FONT_TIMES_NEW_ROMAN, SIZE_FIVE) # Fig num Times
        run_caption_text = p_caption.add_run(caption_text)
        set_run_font(run_caption_text, FONT_KAITI, FONT_KAITI, SIZE_FIVE) # Caption Kaiti
        add_styled_paragraph(doc, "", FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, space_after_pt=6) # Empty line after
        return


    # Add caption
    # "图题采用中文五号楷体，图序号采用 Times New Ramon。图题紧接图的下一行居中打印。"
    p_caption = doc.add_paragraph()
    set_paragraph_formatting(p_caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_25)
    
    run_fig_num = p_caption.add_run(fig_num_text + " ") # e.g., "图3.2 "
    set_run_font(run_fig_num, FONT_KAITI, FONT_TIMES_NEW_ROMAN, SIZE_FIVE) # Fig num Times

    run_caption_text = p_caption.add_run(caption_text) # e.g., "某结构示意图"
    set_run_font(run_caption_text, FONT_KAITI, FONT_KAITI, SIZE_FIVE) # Caption Kaiti

    # "图题后空一行，继续正文内容。"
    add_styled_paragraph(doc, "", FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, space_after_pt=6) # Effectively an empty line with some spacing

def add_table_with_caption(doc, html_table_soup, caption_text_full):
    """Adds a Markdown table (parsed as HTML) and its caption."""
    # Extract table number and caption text
    # E.g., "[表2.1 典型虚拟化环境的物理前缀特征]虚拟化平台"
    # caption_text_full is like "表2.1 典型虚拟化环境的物理前缀特征"
    match = re.match(r'(表\s*[\d\.]+)\s*(.*)', caption_text_full)
    if not match:
        print(f"Warning: Could not parse table caption: {caption_text_full}")
        table_num_text = "表X.X"
        table_title_text = caption_text_full
    else:
        table_num_text = match.group(1).strip()
        table_title_text = match.group(2).strip()

    # Add caption: "表标题采用五号楷体，表序号采用Times New Roman。"居中, 表上方
    p_caption = doc.add_paragraph()
    set_paragraph_formatting(p_caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_25)
    
    run_table_num = p_caption.add_run(table_num_text + " ")
    set_run_font(run_table_num, FONT_KAITI, FONT_TIMES_NEW_ROMAN, SIZE_FIVE) # Table num Times
    
    run_caption_title = p_caption.add_run(table_title_text)
    set_run_font(run_caption_title, FONT_KAITI, FONT_KAITI, SIZE_FIVE) # Table title Kaiti

    # Parse HTML table
    headers = [th.get_text(strip=True) for th in html_table_soup.find_all('th')]
    # Clean the first header if it contained the caption
    if headers and table_title_text in headers[0]: # A bit simplistic
         headers[0] = headers[0].replace(f"[{caption_text_full}]", "").strip()
         if not headers[0]: # If it was only the caption
            # Try to get it from the original markdown structure if possible, or leave as is
            # This part is tricky as the markdown parser might have altered it.
            # For now, we assume the first header cell might be empty or contain the actual header.
            pass


    rows_data = []
    for row_soup in html_table_soup.find('tbody').find_all('tr'):
        cells = [td.get_text(strip=True) for td in row_soup.find_all('td')]
        rows_data.append(cells)

    if not rows_data: # No body rows, maybe it's a header-only table or malformed
        if headers: # If we have headers, create a table with just the header
             table = doc.add_table(rows=1, cols=len(headers))
        else: # No data at all
            print(f"Warning: Table '{caption_text_full}' has no headers or rows. Skipping.")
            return
    else: # Normal case with data rows
        table = doc.add_table(rows=1, cols=len(headers if headers else rows_data[0]))


    table.style = 'Table Grid' # Apply a basic grid style
    table.autofit = True # Allow Word to autofit columns, or set widths manually

    # Populate header row
    if headers:
        for i, header_text in enumerate(headers):
            cell = table.cell(0, i)
            # Remove the caption part from the first header cell text if it's still there
            if i == 0 and f"[{caption_text_full}]" in header_text:
                 header_text = header_text.replace(f"[{caption_text_full}]","").strip()
            
            # Clear existing content and add new run
            cell.text = '' 
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            set_run_font(run, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, bold=True) # Header bold
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER


    # Populate data rows
    for row_idx, row_cells_text in enumerate(rows_data):
        if not headers and row_idx == 0: # If no explicit headers, treat first data row as header
            # This case is less likely with the specified markdown format
            for col_idx, cell_text in enumerate(row_cells_text):
                cell = table.cell(row_idx, col_idx)
                cell.text = ''
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                set_run_font(run, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, bold=True)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            actual_row_idx = row_idx + (1 if headers else 0) # Offset by 1 if headers were added
            if actual_row_idx >= len(table.rows): # Add row if it doesn't exist
                table.add_row()

            for col_idx, cell_text in enumerate(row_cells_text):
                if col_idx >= len(table.columns): # Should not happen if table created correctly
                    print(f"Warning: cell index {col_idx} out of bounds for table {caption_text_full}")
                    continue
                cell = table.cell(actual_row_idx, col_idx)
                cell.text = '' # Clear existing content
                para = cell.paragraphs[0]
                # Add mixed font handling for table cells if needed, for now simple
                run = para.add_run(cell_text)
                set_run_font(run, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT # Or CENTER based on preference for data

    # "表格后空一行，继续正文内容。"
    add_styled_paragraph(doc, "", FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, space_after_pt=6)

def set_default_font_and_line_spacing(doc):
    """Sets the default document font and line spacing for the 'Normal' style."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_TIMES_NEW_ROMAN
    font.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_SONGTI)
    font.size = Pt(SIZE_SMALL_FOUR)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = LINE_SPACING_1_25
    paragraph_format.space_after = Pt(0) # No auto space after for normal paragraphs

def add_page_number_field(paragraph):
    """Adds a PAGE field to a paragraph (typically in footer)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE" #  For just page number. Use "PAGE \* MERGEFORMAT" for more complex.
    run._r.append(instrText)

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar2)

def set_header_footer(section, header_text, footer_font_size_pt):
    """Sets header and footer for a given section."""
    # Header
    header = section.header
    header.is_linked_to_previous = False # Unlink from previous section's header
    p_header = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_header.text = header_text
    set_paragraph_formatting(p_header, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in p_header.runs: # Style existing runs or the new run
        set_run_font(run, FONT_SONGTI, FONT_SONGTI, SIZE_FIVE) # Header宋体五号

    # Add header underline (bottom border to the paragraph)
    pPr = p_header._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single') # Single line
    bottom_border.set(qn('w:sz'), '4')      # Thickness (1/8 pt, so 4 = 0.5pt)
    bottom_border.set(qn('w:space'), '1')   # Space between text and border (1pt)
    bottom_border.set(qn('w:color'), 'auto')# Black
    pBdr.append(bottom_border)
    pPr.append(pBdr)
    
    # Footer (for page numbers)
    footer = section.footer
    footer.is_linked_to_previous = False # Unlink
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.clear() # Clear any existing content
    set_paragraph_formatting(p_footer, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Add page number field
    add_page_number_field(p_footer)
    # Style the (potential) runs of the page number field - this is tricky as field is dynamic
    # We set the paragraph font, Word should apply it.
    for run in p_footer.runs: # Style existing runs
        set_run_font(run, FONT_SONGTI, FONT_SONGTI, footer_font_size_pt)
    # If no runs, set font on paragraph level (less direct for fields)
    if not p_footer.runs:
        # This is a bit of a hack for fields; ideally Word picks up the paragraph's default.
        # We add a dummy run to set font, then clear it.
        # Or, rely on Word to format the PAGE field correctly.
        # For now, let's assume Word handles it based on paragraph's implied font.
        # To be safe, one might need to inspect and style runs *after* Word updates fields.
        # For script generation, setting paragraph font is the best guess.
        pf = p_footer.paragraph_format
        # This doesn't directly set font for the field, but it's a hint.
        # The run added by add_page_number_field should ideally inherit.

def add_toc_placeholder(doc):
    """Adds a paragraph indicating where the TOC should be, and a TOC field."""
    # TOC Title
    toc_title_p = doc.add_paragraph()
    set_paragraph_formatting(toc_title_p, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_val=LINE_SPACING_1_5, space_before_pt=12, space_after_pt=12)
    run_toc_title = toc_title_p.add_run("目   录") # Spaced as per some conventions
    set_run_font(run_toc_title, FONT_HEITI, FONT_HEITI, SIZE_THREE, bold=True) # Assuming Heiti Three for TOC title

    # TOC Field (Word will populate this)
    # \o "1-3" includes heading levels 1 to 3
    # \h creates hyperlinks
    # \z hides tab leader and page number in Web layout view
    # \u uses outline levels from paragraphs
    toc_field_p = doc.add_paragraph()
    run = toc_field_p.add_run()
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u' # Standard TOC field
    
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate) # Optional: if you want to show something before update
    # Example: run._r.append(parse_xml('<w:t>Right-click to update field.</w:t>'))
    run._r.append(fldChar_end)
    
    # Default formatting for TOC entries (Word will apply its own styles based on Heading X styles)
    # User should ensure their Word's Heading 1, 2, 3 styles are appropriate or modify them.
    # The script primarily ensures the source headings are marked correctly.
    # Set TOC paragraph to a common font, Word styles will override for entries.
    set_paragraph_formatting(toc_field_p, line_spacing_val=LINE_SPACING_1_25)
    set_run_font(run, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR)

ReferencesList = []

def preprocess_html(soup):
    global ReferencesList
    for sup in soup.find_all('sup', id=re.compile(r'^fnref:\d+$')):
        a = sup.find('a', class_='footnote-ref')
        if a and a.string and a.string.strip().isdigit():
            footnote_number = a.string.strip()
            sup.replace_with(f'[^{footnote_number}]')
    
    for li in soup.select('div.footnote ol li'):
        p = li.find('p')
        if p:
            for a in p.find_all('a'):
                a.decompose()
            text = p.get_text(strip=True)
            ReferencesList.append(text)
    return soup

# --- Main Processing Logic ---
def markdown_to_word(md_content, output_docx_path, temp_dir_for_mermaid):
    doc = Document()

    # 1. Page Setup (A4, Margins) for the first default section
    section = doc.sections[0]
    section.page_height = Cm(29.7) # A4 height
    section.page_width = Cm(21.0)  # A4 width
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)

    # 2. Set Default Font and Line Spacing for 'Normal' style
    set_default_font_and_line_spacing(doc)

    # 3. Parse Markdown to HTML
    # 'extra' includes tables, fenced_code, footnotes, etc.
    # 'sane_lists' for better list handling.
    # 'nl2br' if you want newlines in markdown to become <br> (thesis usually doesn't want this for paragraphs)
    html_content = md_parser.markdown(md_content, extensions=['extra', 'footnotes', 'tables', 'sane_lists', 'meta', 'toc'])
    soup = BeautifulSoup(html_content, 'html.parser')
    soup = preprocess_html(soup)

    # State variables
    current_chapter_num_str = ""
    processed_elements_count = 0
    main_text_started = False
    toc_items = [] # For a manually generated TOC if needed, or just for tracking

    # --- Iterate through Markdown elements (now HTML) ---
    elements = list(soup.children)
    idx = 0
    while idx < len(elements):
        el = elements[idx]
        idx += 1
        processed_elements_count +=1
        
        if el.name is None: # NavigableString, likely whitespace
            continue

        # --- Section Titles (Abstracts, Main Chapters, Refs, Ack) ---
        if el.name == 'h1':
            title_text = el.get_text(strip=True)
            
            if "摘要" == title_text and not "ABSTRACT" in title_text: # Chinese Abstract
                # Title "摘 要"
                p_title = doc.add_paragraph()
                set_paragraph_formatting(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=Pt(SIZE_THREE)) # Space after title
                run = p_title.add_run("摘   要") # Spaced
                set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_THREE, bold=True)
                
                # Content:楷体四号,行距固定值20磅,段首缩进
                # Next element should be the abstract content
                while idx < len(elements) and elements[idx].name != 'h1':
                    if elements[idx].name == 'p':
                        abstract_content = elements[idx].get_text(strip=True)
                        #print(abstract_content)
                        add_styled_paragraph(doc, abstract_content, FONT_KAITI, FONT_KAITI, SIZE_FOUR, # Kaiti for all
                                         first_line_indent_cm=0.7, # Approx 2 chars
                                         fixed_line_height_pt=LINE_SPACING_FIXED_20PT,
                                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) # Or LEFT
                    idx += 1 # Consume abstract paragraph
                doc.add_paragraph() # Blank line after Chinese Abstract content (as per "空一行") - or use space_after_pt

            elif "ABSTRACT" == title_text: # English Abstract
                # "空两行后" from Chinese abstract - effectively one blank para + space_before on ABSTRACT
                add_styled_paragraph(doc, "", FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR, space_after_pt=Pt(SIZE_SMALL_FOUR)) # One blank line

                # English Title and Subtitle (assuming they are H2/H3 or paras before ABSTRACT H1 in MD)
                # This script assumes ABSTRACT H1 is the main marker. User needs to ensure MD structure.
                # For simplicity, we're not parsing separate Eng title/subtitle from MD here.

                # "ABSTRACT" heading
                p_title = doc.add_paragraph()
                set_paragraph_formatting(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=Pt(SIZE_THREE))
                run = p_title.add_run("ABSTRACT")
                set_run_font(run, FONT_HEITI, FONT_TIMES_NEW_ROMAN, SIZE_THREE, bold=True)

                idx += 1
                # Content: 四号, 段首缩进, 行距固定值20磅
                while idx < len(elements) and elements[idx].name != 'h1':
                    if elements[idx].name == 'p':
                        abstract_content = elements[idx].get_text(strip=True)
                        #print(abstract_content)
                        add_styled_paragraph(doc, abstract_content, FONT_TIMES_NEW_ROMAN, FONT_TIMES_NEW_ROMAN, SIZE_FOUR,
                                         first_line_indent_cm=0.7, # Approx 2 chars
                                         fixed_line_height_pt=LINE_SPACING_FIXED_20PT,
                                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) # Or LEFT
                    idx += 1
                
                # --- Add Table of Contents after English Abstract ---
                add_toc_placeholder(doc)
                
                # --- Start new section for main text for headers/footers ---
                # "页眉从正文开始处起到致谢结束处终止"
                # This means Abstracts and TOC are in section 0, main text starts section 1
                doc.add_section(WD_SECTION_START.NEW_PAGE)
                main_text_started = True
                # Apply header/footer to this new section and subsequent ones
                # (Loop through sections later to apply)

            elif title_text.startswith("第") and "章" in title_text: # Main Text Chapter
                if not main_text_started: # If somehow first chapter appears before ABSTRACT is done
                    doc.add_section(WD_SECTION_START.NEW_PAGE)
                    main_text_started = True

                chap_match = re.match(r"(第[一二三四五六七八九十百]+章)\s*(.*)", title_text)
                if chap_match:
                    current_chapter_num_str = chap_match.group(1) # e.g. "第一章"
                    # The add_heading function handles the full "第X章 XXX" format
                    _, toc_entry_text = add_heading(doc, title_text, 1)
                    toc_items.append({'level': 1, 'text': toc_entry_text, 'page': '?'})
                else: # Fallback if format is slightly off
                     _, toc_entry_text = add_heading(doc, title_text, 1) # Fallback
                     toc_items.append({'level': 1, 'text': toc_entry_text, 'page': '?'})


            elif "参考文献" == title_text:
                if not main_text_started: # Should not happen
                    doc.add_section(WD_SECTION_START.NEW_PAGE)
                    main_text_started = True
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE) # New page

                '''
                p_title = doc.add_paragraph()
                set_paragraph_formatting(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=Pt(SIZE_SMALL_FOUR))
                run = p_title.add_run("参考文献")
                set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_THREE, bold=True)
                # References are handled by footnote processing later or specific list items
                '''
                _, toc_entry_text = add_heading(doc, title_text, 1)
                toc_items.append({'level': 1, 'text': toc_entry_text, 'page': '?'})

                for ref_idx, ref in enumerate(ReferencesList):
                    ref_text = f"[{ref_idx + 1}] {ref}"
                    #print(ref_text)
                    add_styled_paragraph(doc, ref_text, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                         line_spacing_val=LINE_SPACING_1_25, first_line_indent_cm=0.7, is_reference=True)

            elif "致谢" == title_text:
                if not main_text_started: # Should not happen
                    doc.add_section(WD_SECTION_START.NEW_PAGE)
                    main_text_started = True
                # No new page needed if it follows refs unless specified, but often is.
                # The prompt doesn't explicitly say new page for 致谢 if after refs.
                # Let's assume it continues unless it's the first thing in a new section.
                '''
                p_title = doc.add_paragraph()
                set_paragraph_formatting(p_title, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before_pt=Pt(SIZE_THREE), space_after_pt=Pt(SIZE_SMALL_FOUR))
                run = p_title.add_run("致       谢") # "致"与"谢"之间空四格 (two Chinese chars)
                set_run_font(run, FONT_HEITI, FONT_HEITI, SIZE_THREE, bold=True)
                '''
                _, toc_entry_text = add_heading(doc, title_text, 1)
                toc_items.append({'level': 1, 'text': toc_entry_text, 'page': '?'})
                # Content: 小四号宋体, 1.25倍行距
                # Subsequent paragraphs will be handled by the 'p' tag logic.

        # --- Main Text Headings (H2, H3, H4, H5) ---
        elif el.name == 'h2':
            _, toc_entry_text = add_heading(doc, el.get_text(strip=True), 2)
            toc_items.append({'level': 2, 'text': toc_entry_text, 'page': '?'})
        elif el.name == 'h3':
            _, toc_entry_text = add_heading(doc, el.get_text(strip=True), 3)
            toc_items.append({'level': 3, 'text': toc_entry_text, 'page': '?'})
        elif el.name == 'h4': # 小节内的小标题序号用1、2、3……，小标题用黑体字单列一行
            add_heading(doc, el.get_text(strip=True), 4) # Handled by add_heading
        elif el.name == 'h5': # 其余层次序号用⑴、⑵、⑶……
             # These are not true "headings" in Word style, but styled paragraphs.
             # Markdown H5 -> styled paragraph
            add_styled_paragraph(doc, el.get_text(strip=True), FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                 line_spacing_val=LINE_SPACING_1_25)


        # --- Paragraphs (Default text, lists, blockquotes) ---
        elif el.name == 'p':
            # Check for images within paragraphs
            img_tag = el.find('img')
            if img_tag:
                # Format: ![图2.1 某结构示意图](图片地址)
                # alt="图2.1 某结构示意图", src="图片地址"
                alt_text = img_tag.get('alt', '')
                img_src = img_tag.get('src', '')

                caption_match = re.match(r'(图\s*[\d\.]+)\s*(.*)', alt_text)
                if caption_match:
                    fig_num_text = caption_match.group(1).strip()
                    caption_content = caption_match.group(2).strip()
                    add_image_with_caption(doc, img_src, caption_content, fig_num_text)
                else:
                    # Fallback if alt text is not in the specified "图X.Y Caption" format
                    add_image_with_caption(doc, img_src, alt_text if alt_text else "Untitled Image", "图?.?")
            else:
                # Regular paragraph
                para_text = el.get_text(strip=True) # Get raw text
                if para_text: # Avoid adding empty paragraphs unless intended for spacing
                    # Check if this paragraph is part of "致谢" content
                    # This is a heuristic; better to rely on structure if "致谢" H1 is followed by <p>
                    # For now, assume default paragraph styling.
                    # Specific styling for "致谢" content paragraphs:
                    # Check if the previous H1 was "致谢" - needs more state tracking or lookbehind in soup.
                    # Simplified: if in a section after "致谢" H1, apply.
                    # This needs refinement if other H1s can appear after "致谢".

                    # Default paragraph handling
                    add_styled_paragraph(doc, para_text, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                         line_spacing_val=LINE_SPACING_1_25, first_line_indent_cm=0.7) # Default indent for body
        
        # --- Lists (ul, ol) for References or general content ---
        elif el.name in ['ul', 'ol']:
            # This is where Footnotes from markdown `[^N]: text` often end up as list items.
            # Or, they could be regular lists.
            for item_idx, li in enumerate(el.find_all('li', recursive=False)):
                item_text = li.get_text(strip=True)
                # Check if it's a reference item (markdown footnote style)
                # `[^1]: ayoubfaouzi...` becomes `[1] ayoubfaouzi...`
                ref_match = re.match(r'\[\^(\d+)\]:\s*(.*)', item_text, re.DOTALL) # Original footnote def
                if ref_match: # This is a bibliography item
                    ref_num = ref_match.group(1)
                    ref_content = ref_match.group(2).strip()
                    
                    # "参考文献" "小四号宋体、左起、悬挂缩进、1.25倍行距"
                    p_ref = add_styled_paragraph(doc, f"[{ref_num}] {ref_content}", 
                                             FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                             line_spacing_val=LINE_SPACING_1_25,
                                             alignment=WD_ALIGN_PARAGRAPH.LEFT)
                    # Apply hanging indent
                    # Hanging indent: first_line_indent is negative, left_indent is positive
                    # python-docx uses first_line_indent for hanging if negative.
                    # A common hanging indent is 0.5 inches or 1.27 cm.
                    # Let's use a value that aligns with typical [1] numbering.
                    p_ref.paragraph_format.left_indent = Cm(0.7) # Indent entire paragraph
                    p_ref.paragraph_format.first_line_indent = Cm(-0.7) # Hang the first line back
                else:
                    # Regular list item
                    prefix = f"{item_idx + 1}. " if el.name == 'ol' else "- "
                    add_styled_paragraph(doc, prefix + item_text, FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                         line_spacing_val=LINE_SPACING_1_25, first_line_indent_cm=0.7,
                                         space_before_pt=2, space_after_pt=2)


        # --- Tables ---
        elif el.name == 'table':
            # Markdown: | [表2.1 Caption]Header1 | Header2 |
            # Caption is in the first header cell.
            first_th = el.find('th')
            caption_text_full = ""
            if first_th:
                caption_match = re.search(r'\[(表.*?)]', first_th.get_text(strip=True))
                if caption_match:
                    caption_text_full = caption_match.group(1)
            
            if caption_text_full:
                add_table_with_caption(doc, el, caption_text_full)
            else:
                print(f"Warning: Table found without a recognized caption format: {el.prettify()[:100]}")
                # Add table without caption or with a placeholder
                add_table_with_caption(doc, el, "表?.? Unknown Table")


        # --- Fenced Code Blocks (Mermaid or other code) ---
        elif el.name == 'pre':
            code_tag = el.find('code')
            ##code_tag = pre_tag.find('code') if pre_tag else None
            
            if code_tag:
                code_content = code_tag.get_text() # Keep original newlines
                # Check for Mermaid: ```mermaid ... ```
                # The 'language-mermaid' class might be on `code` or `pre`
                is_mermaid = 'language-mermaid' in code_tag.get('class', [])

                if is_mermaid:
                    # Extract caption: %%图3.1 某功能流程图
                    mermaid_caption_match = re.match(r'%%(图\s*[\d\.]+)\s*(.*)', code_content) # From first line
                    if mermaid_caption_match:
                        fig_num_text = mermaid_caption_match.group(1).strip()
                        caption_content = mermaid_caption_match.group(2).strip()
                        
                        # Remove the caption line from the mermaid code itself
                        mermaid_code_for_render = '\n'.join(code_content.splitlines()[1:])
                        
                        #print(f'idx = {idx}')
                        image_file = convert_mermaid_to_image_mmdc(mermaid_code_for_render, temp_dir_for_mermaid)
                        if image_file:
                            add_image_with_caption(doc, image_file, caption_content, fig_num_text)
                            os.remove(image_file) # Clean up temp image
                        else:
                            print(f"Failed to convert Mermaid diagram: {fig_num_text} {caption_content}")
                            # Add placeholder text for failed mermaid
                            add_styled_paragraph(doc, f"[Failed to render Mermaid diagram: {fig_num_text} {caption_content}]",
                                                 FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                                 line_spacing_val=LINE_SPACING_1_25)
                    else:
                        print(f"Warning: Mermaid diagram found without '%%图X.Y Caption' comment: {code_content[:50]}")
                        # Add placeholder for mermaid without caption
                        add_styled_paragraph(doc, "[Mermaid diagram - caption not found]",
                                             FONT_SONGTI, FONT_TIMES_NEW_ROMAN, SIZE_SMALL_FOUR,
                                             line_spacing_val=LINE_SPACING_1_25)

                else:
                    # Regular code block - not typically part of a thesis body, or needs specific formatting
                    # For now, just add as preformatted text (simple paragraph, monospace font)
                    p_code = add_styled_paragraph(doc, code_content, "Courier New", "Courier New", SIZE_SMALL_FOUR - 2, # Smaller for code
                                             line_spacing_val=1.0, space_before_pt=5, space_after_pt=5)
                    p_code.paragraph_format.left_indent = Cm(1)
                    # Could add a border or background shading if desired.
        elif el.name == 'hr':
            # Horizontal rule, could be a page break or section break in some contexts
            # For now, just add a paragraph with a symbolic representation or ignore
            # doc.add_paragraph("--- horizontal rule ---")
            pass # Often not needed in final doc

    # --- Apply Headers and Footers ---
    # Header: custom header, 宋体五号居中, 加页眉线
    # Page numbers: 宋体小五号, 居中
    # Scope: From main text start (section 1 onwards) to acknowledgments end.
    header_text_content = config['Artical-Header']
    for i, section_to_modify in enumerate(doc.sections):
        if i == 0: # Section 0 (Abstracts, TOC) - no header/footer or different
            section_to_modify.header.is_linked_to_previous = True # Link to default empty or make it distinct
            section_to_modify.footer.is_linked_to_previous = True
            # Clear any accidental header/footer from section 0
            if section_to_modify.header.paragraphs:
                for p in section_to_modify.header.paragraphs:
                    p.clear()
            if section_to_modify.footer.paragraphs:
                 for p in section_to_modify.footer.paragraphs:
                    p.clear()
        else: # Sections from main text onwards
            set_header_footer(section_to_modify, header_text_content, SIZE_SMALL_FIVE)


    # --- Save Document ---
    doc.save(output_docx_path)
    print(f"Document saved to {output_docx_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert Markdown to formatted Word document for thesis.")
    parser.add_argument("md_file", help="Path to the input Markdown file.")
    parser.add_argument("docx_file", help="Path to the output Word DOCX file.")
    args = parser.parse_args()

    if not os.path.exists(args.md_file):
        print(f"Error: Markdown file not found: {args.md_file}")
        exit(1)

    with open(args.md_file, 'r', encoding='utf-8') as f:
        md_content_main = f.read()

    # Create a temporary directory for Mermaid image conversion
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            markdown_to_word(md_content_main, args.docx_file, temp_dir)
        except Exception as e:
            print(f"An unexpected error occurred: {e}")
            import traceback
            traceback.print_exc()